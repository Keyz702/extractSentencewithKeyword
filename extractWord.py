from typing_extensions import final
import docx
import pandas as pd

templ = docx.Document('filename')
templ_tables = templ.tables

type = type(templ.paragraphs)

filePath = 'results-filepath'

def getKeySentence(a):
    final_data = []
    for i in a.paragraphs:
        raw_data = i.text
        out_data = [sentence + '.' for sentence in raw_data.split('.') if 'keyword' in sentence]
        clean_data = []
        for item in out_data:
            if item!=[]:
                clean_data.append(item)
        for x in clean_data:
            final_data.append(x)
    return final_data

section = templ.sections[0]
header = section.header
for paragraph in header.paragraph:
    print(paragraph.text)

dataframe = pd.DataFrame(getKeySentence(templ))
dataframe.to_csv(filePath, header=False, index=False)

